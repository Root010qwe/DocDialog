import logging
from pathlib import Path

from app.parsers.base import AbstractParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(AbstractParser):

    EXTENSIONS = {".docx", ".doc"}
    CONTENT_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def can_handle(self, content_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return content_type in self.CONTENT_TYPES or ext in self.EXTENSIONS

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        title = Path(filename).stem
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)

            parts: list[str] = []

            if doc.core_properties.title:
                title = doc.core_properties.title

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))

            text = "\n\n".join(parts)
            logger.info("DocxParser extracted %d chars from %s", len(text), filename)
            return ParsedDocument(text=text, title=title)

        except Exception as e:
            logger.error("DocxParser failed for %s: %s", filename, e)
            return ParsedDocument(text="", title=title)
