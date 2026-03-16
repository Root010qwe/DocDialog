import logging
from pathlib import Path

from app.parsers.base import AbstractParser, ParsedDocument

logger = logging.getLogger(__name__)

DOCLING_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/html",
    "application/xhtml+xml",
}

DOCLING_EXTENSIONS = {".pdf", ".docx", ".doc", ".html", ".htm"}


class DoclingParser(AbstractParser):
    """Uses Docling to extract structured text from PDF, DOCX, HTML."""

    def can_handle(self, content_type: str, filename: str) -> bool:
        ext = Path(filename).suffix.lower()
        return content_type in DOCLING_TYPES or ext in DOCLING_EXTENSIONS

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(file_path)
            doc = result.document

            # Export to markdown for clean text extraction
            text = doc.export_to_markdown()
            title = Path(filename).stem

            return ParsedDocument(text=text, title=title)

        except Exception as e:
            logger.warning("Docling failed for %s: %s — falling back", filename, e)
            return self._fallback(file_path, filename)

    def _fallback(self, file_path: str, filename: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        title = Path(filename).stem
        try:
            if ext == ".pdf":
                import pdfplumber
                texts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            texts.append(t)
                return ParsedDocument(text="\n\n".join(texts), title=title)

            elif ext in (".docx", ".doc"):
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return ParsedDocument(text="\n\n".join(paragraphs), title=title)

            elif ext in (".html", ".htm"):
                from bs4 import BeautifulSoup
                with open(file_path, encoding="utf-8", errors="ignore") as f:
                    soup = BeautifulSoup(f.read(), "lxml")
                return ParsedDocument(text=soup.get_text(separator="\n"), title=title)

        except Exception as e2:
            logger.error("Fallback parser also failed for %s: %s", filename, e2)

        return ParsedDocument(text="", title=title)
