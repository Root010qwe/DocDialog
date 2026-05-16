import io
import uuid
from datetime import datetime, timezone

from sqlalchemy.ext.asyncio import AsyncSession

from app.services.dialog_service import DialogService


class ExportService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.dialog_service = DialogService(session)

    @staticmethod
    def _register_pdf_font():
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os

        if "DejaVuSans" in pdfmetrics.getRegisteredFontNames():
            return
        candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/DejaVuSans.ttf",
        ]
        for path in candidates:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont("DejaVuSans", path))
                bold_path = path.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf")
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", bold_path))
                return

    async def export_dialog_pdf(self, dialog_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
        )

        self._register_pdf_font()
        font = "DejaVuSans"

        dialog = await self.dialog_service.get_dialog(dialog_id, user_id)
        if not dialog:
            return b""

        messages = await self.dialog_service.get_messages_with_citations(dialog_id)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        style_title = ParagraphStyle(
            "DocTitle",
            parent=styles["Normal"],
            fontName=font,
            fontSize=18,
            spaceAfter=6,
            textColor=colors.HexColor("#1a1a2e"),
        )
        style_meta = ParagraphStyle(
            "DocMeta",
            parent=styles["Normal"],
            fontName=font,
            fontSize=9,
            textColor=colors.HexColor("#666666"),
            spaceAfter=4,
        )
        style_user = ParagraphStyle(
            "DocUser",
            parent=styles["Normal"],
            fontName=font,
            fontSize=11,
            leftIndent=0,
            backColor=colors.HexColor("#e8f4fd"),
            borderPad=8,
            spaceAfter=6,
        )
        style_assistant = ParagraphStyle(
            "DocAssistant",
            parent=styles["Normal"],
            fontName=font,
            fontSize=11,
            leftIndent=0,
            backColor=colors.HexColor("#f0f7f0"),
            borderPad=8,
            spaceAfter=6,
        )
        style_role = ParagraphStyle(
            "DocRole",
            parent=styles["Normal"],
            fontName=font,
            fontSize=9,
            textColor=colors.HexColor("#888888"),
            spaceBefore=8,
            spaceAfter=2,
        )
        style_citation = ParagraphStyle(
            "DocCitation",
            parent=styles["Normal"],
            fontName=font,
            fontSize=9,
            leftIndent=1 * cm,
            textColor=colors.HexColor("#555555"),
            spaceAfter=2,
        )

        story = []

        title_text = dialog.title or "Диалог без названия"
        story.append(Paragraph(f"DocDialog — {title_text}", style_title))

        created = dialog.created_at
        if hasattr(created, "strftime"):
            date_str = created.strftime("%d.%m.%Y %H:%M")
        else:
            date_str = str(created)
        story.append(Paragraph(f"Создан: {date_str}   |   Сообщений: {len(messages)}", style_meta))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dddddd")))
        story.append(Spacer(1, 0.3 * cm))

        for msg in messages:
            role_label = "Пользователь" if msg.role.value == "user" else "Ассистент"
            msg_time = msg.created_at.strftime("%H:%M") if hasattr(msg.created_at, "strftime") else ""
            story.append(Paragraph(f"{role_label}  {msg_time}", style_role))

            style = style_user if msg.role.value == "user" else style_assistant
            content_safe = (msg.content or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(content_safe, style))

            citations = getattr(msg, "citations", None) or []
            if citations:
                story.append(Paragraph("Источники:", style_citation))
                for i, cit in enumerate(citations[:5], 1):
                    doc_name = cit.get("document_title") or cit.get("document_name", "Документ")
                    chunk_text = (cit.get("chunk_text") or cit.get("text", ""))[:150]
                    score = cit.get("similarity_score", 0)
                    cit_text = f"{i}. <b>{doc_name}</b> (релевантность: {score:.2f}) — {chunk_text}…"
                    cit_safe = cit_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    # Restore bold tags after escaping
                    cit_safe = cit_safe.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
                    story.append(Paragraph(cit_safe, style_citation))

            story.append(Spacer(1, 0.2 * cm))

        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#dddddd")))
        now = datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M UTC")
        story.append(Paragraph(f"Экспортировано: {now} | DocDialog — RAG система диалога с документами", style_meta))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    async def export_dialog_docx(self, dialog_id: uuid.UUID, user_id: uuid.UUID) -> bytes:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm

        dialog = await self.dialog_service.get_dialog(dialog_id, user_id)
        if not dialog:
            return b""

        messages = await self.dialog_service.get_messages_with_citations(dialog_id)

        docx = DocxDocument()

        for section in docx.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

        title_para = docx.add_heading(level=1)
        title_run = title_para.add_run(f"DocDialog — {dialog.title or 'Диалог без названия'}")
        title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        created = dialog.created_at
        date_str = created.strftime("%d.%m.%Y %H:%M") if hasattr(created, "strftime") else str(created)
        meta = docx.add_paragraph()
        meta_run = meta.add_run(f"Создан: {date_str}   |   Сообщений: {len(messages)}")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        docx.add_paragraph()

        for msg in messages:
            role_label = "Пользователь" if msg.role.value == "user" else "Ассистент"
            msg_time = msg.created_at.strftime("%H:%M") if hasattr(msg.created_at, "strftime") else ""
            role_para = docx.add_paragraph()
            role_run = role_para.add_run(f"{role_label}  {msg_time}")
            role_run.font.size = Pt(9)
            role_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            role_run.bold = True

            content_para = docx.add_paragraph()
            content_run = content_para.add_run(msg.content or "")
            content_run.font.size = Pt(11)
            if msg.role.value == "user":
                content_run.font.color.rgb = RGBColor(0x1A, 0x5F, 0xA3)
            else:
                content_run.font.color.rgb = RGBColor(0x1E, 0x6B, 0x3A)

            citations = getattr(msg, "citations", None) or []
            if citations:
                cit_header = docx.add_paragraph()
                cit_h_run = cit_header.add_run("Источники:")
                cit_h_run.font.size = Pt(9)
                cit_h_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                for cit in citations[:5]:
                    doc_name = cit.get("document_title") or cit.get("document_name", "Документ")
                    chunk_text = (cit.get("chunk_text") or cit.get("text", ""))[:150]
                    score = cit.get("similarity_score", 0)
                    cit_para = docx.add_paragraph(style="List Bullet")
                    cit_run = cit_para.add_run(f"{doc_name} (релевантность: {score:.2f}) — {chunk_text}…")
                    cit_run.font.size = Pt(9)
                    cit_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            docx.add_paragraph()

        footer = docx.add_paragraph()
        now = datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M UTC")
        footer_run = footer.add_run(
            f"Экспортировано: {now} | DocDialog — RAG система диалога с документами"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        buffer = io.BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        return buffer.read()
